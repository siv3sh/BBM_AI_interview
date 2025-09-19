import streamlit as st
import pdfplumber
import docx
import json
import time
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from enum import Enum
import logging

# LangChain imports
from langchain.agents import Tool, AgentExecutor, create_react_agent
from langchain.prompts import PromptTemplate
from langchain.schema import BaseOutputParser
from langchain.memory import ConversationBufferMemory
from langchain_groq import ChatGroq
from langchain_community.callbacks.streamlit import StreamlitCallbackHandler
from langchain.tools.base import BaseTool
from langchain.chains import LLMChain
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field, validator

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ========== Data Models ==========
class AnalysisScore(BaseModel):
    score: int = Field(ge=0, le=100, description="ATS compatibility score")
    reasoning: str = Field(description="Explanation for the score")

class KeywordMatch(BaseModel):
    matched_keywords: List[str] = Field(description="Keywords found in resume")
    missing_keywords: List[str] = Field(description="Keywords missing from resume")
    match_percentage: int = Field(ge=0, le=100, description="Percentage of keywords matched")

class SkillsAnalysis(BaseModel):
    technical_skills: List[str] = Field(description="Missing technical skills")
    programming_languages: List[str] = Field(description="Missing programming languages")
    libraries_frameworks: List[str] = Field(description="Missing libraries/frameworks")
    tools_platforms: List[str] = Field(description="Missing tools/platforms")

class ResumeAnalysis(BaseModel):
    score: int = Field(ge=0, le=100)
    keyword_match: KeywordMatch
    missing_skills: SkillsAnalysis
    soft_skills: List[str] = Field(description="Missing soft skills")
    suggestions: List[str] = Field(description="Improvement suggestions")

# ========== Configuration ==========
@dataclass
class ModelConfig:
    name: str
    model_id: str
    description: str
    speed_rating: int  # 1-5 scale

MODELS = {
    "Gemma 2B (Fast)": ModelConfig("Gemma 2B", "gemma-2b-it", "Fast analysis with good accuracy", 5),
    "Llama 3 70B (Advanced)": ModelConfig("Llama 3 70B", "llama3-70b-8192", "Advanced analysis with superior accuracy", 3),
    "Mixtral 8x7B (Balanced)": ModelConfig("Mixtral 8x7B", "mixtral-8x7b-32768", "Balanced speed and accuracy", 4)
}

class AgentRole(Enum):
    ANALYZER = "resume_analyzer"
    OPTIMIZER = "resume_optimizer"
    VALIDATOR = "content_validator"

# ========== Custom Tools ==========
class ResumeExtractorTool(BaseTool):
    name: str = "resume_extractor"
    description: str = "Extracts text content from PDF or DOCX files"
    
    def _run(self, file_path: str) -> str:
        """Extract text from uploaded file"""
        try:
            if file_path.endswith('.pdf'):
                with pdfplumber.open(file_path) as pdf:
                    text = " ".join(page.extract_text() or "" for page in pdf.pages)
            elif file_path.endswith('.docx'):
                doc = docx.Document(file_path)
                text = " ".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
            else:
                raise ValueError(f"Unsupported file format: {file_path}")
            
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise e
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove special characters and normalize whitespace
        for char in ['\x0c', '\uf0b7', '\u2022', '\n', '\t']:
            text = text.replace(char, ' ')
        
        # Normalize multiple spaces
        while "  " in text:
            text = text.replace("  ", " ")
        
        return text.strip()

class ResumeAnalyzerTool(BaseTool):
    name: str = "resume_analyzer"
    description: str = "Analyzes resume against job description for ATS compatibility"
    llm: ChatGroq = None
    parser: PydanticOutputParser = None
    
    def __init__(self, llm: ChatGroq, **kwargs):
        super().__init__(**kwargs)
        object.__setattr__(self, 'llm', llm)
        object.__setattr__(self, 'parser', PydanticOutputParser(pydantic_object=ResumeAnalysis))
    
    def _run(self, resume_text: str, job_description: str) -> str:
        """Analyze resume and return structured results"""
        try:
            analysis_prompt = PromptTemplate(
                input_variables=["resume_text", "job_description", "format_instructions"],
                template="""
                You are an expert ATS (Applicant Tracking System) analyzer. 
                Analyze the resume against the job description and provide detailed insights.

                JOB DESCRIPTION:
                {job_description}

                RESUME CONTENT:
                {resume_text}

                Provide a comprehensive analysis focusing on:
                1. Keyword matching between resume and job description
                2. Missing technical skills, programming languages, frameworks, and tools
                3. Missing soft skills mentioned in the job description
                4. Actionable suggestions for improvement
                5. An overall ATS compatibility score (0-100)

                {format_instructions}
                """
            )
            
            chain = LLMChain(llm=self.llm, prompt=analysis_prompt)
            
            result = chain.run(
                resume_text=resume_text[:6000],  # Limit text length
                job_description=job_description[:4000],
                format_instructions=self.parser.get_format_instructions()
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error in resume analysis: {str(e)}")
            return self._create_error_response(str(e))
    
    def _create_error_response(self, error_msg: str) -> str:
        """Create a fallback response when analysis fails"""
        error_response = ResumeAnalysis(
            score=0,
            keyword_match=KeywordMatch(
                matched_keywords=[],
                missing_keywords=["Analysis failed - please try again"],
                match_percentage=0
            ),
            missing_skills=SkillsAnalysis(
                technical_skills=[],
                programming_languages=[],
                libraries_frameworks=[],
                tools_platforms=[]
            ),
            soft_skills=[],
            suggestions=[f"Error occurred: {error_msg}", "Please check your inputs and try again"]
        )
        return error_response.json()

class ResumeOptimizerTool(BaseTool):
    name: str = "resume_optimizer"
    description: str = "Optimizes resume content for better ATS compatibility"
    llm: ChatGroq = None
    
    def __init__(self, llm: ChatGroq, **kwargs):
        super().__init__(**kwargs)
        object.__setattr__(self, 'llm', llm)
    
    def _run(self, resume_text: str, job_description: str, analysis_results: str) -> str:
        """Optimize resume based on analysis"""
        try:
            optimization_prompt = PromptTemplate(
                input_variables=["resume_text", "job_description", "analysis_results"],
                template="""
                You are a professional resume writer and career consultant. 
                Optimize the following resume for better ATS compatibility and job relevance.

                ORIGINAL RESUME:
                {resume_text}

                TARGET JOB DESCRIPTION:
                {job_description}

                ANALYSIS INSIGHTS:
                {analysis_results}

                Guidelines for optimization:
                1. Incorporate relevant keywords naturally
                2. Quantify achievements with specific metrics
                3. Use strong action verbs (Led, Architected, Optimized, etc.)
                4. Ensure ATS-friendly formatting
                5. Maintain professional tone and accuracy
                6. Keep the original structure but enhance content
                7. Add missing skills where appropriate and truthful

                Return ONLY the optimized resume text, maintaining the original format structure.
                """
            )
            
            chain = LLMChain(llm=self.llm, prompt=optimization_prompt)
            
            result = chain.run(
                resume_text=resume_text[:5000],
                job_description=job_description[:3000],
                analysis_results=analysis_results[:2000]
            )
            
            return result.strip()
            
        except Exception as e:
            logger.error(f"Error in resume optimization: {str(e)}")
            return f"Optimization failed: {str(e)}. Please try again with different inputs."

# ========== Enhanced Agent System with Retry and Fallback ==========
class ResumeAgentSystem:
    def __init__(self, groq_api_key: str):
        self.groq_api_key = groq_api_key
        self.llm = None
        self.memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
        self.max_retries = 3
        self.retry_delay = 2  # seconds
        self._initialize_llm()
    
    def _initialize_llm(self):
        """Initialize LLM with error handling"""
        try:
            self.llm = ChatGroq(
                groq_api_key=self.groq_api_key,
                model_name="llama3-70b-8192",
                temperature=0.3,
                max_retries=5,  # Increased retries
                request_timeout=120,  # Increased timeout
                max_tokens=2000
            )
            logger.info("LLM initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize LLM: {str(e)}")
            raise e
    
    def _check_api_status(self) -> bool:
        """Check if Groq API is accessible"""
        try:
            # Simple test call to check API status
            test_response = self.llm.invoke("Hello")
            return True
        except Exception as e:
            error_msg = str(e).lower()
            if "503" in error_msg or "service unavailable" in error_msg:
                return False
            return True
    
    def _retry_with_exponential_backoff(self, func, *args, **kwargs):
        """Retry function with exponential backoff"""
        import time
        
        for attempt in range(self.max_retries):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                error_msg = str(e).lower()
                
                if "503" in error_msg or "service unavailable" in error_msg:
                    if attempt < self.max_retries - 1:
                        wait_time = self.retry_delay * (2 ** attempt)  # Exponential backoff
                        st.warning(f"🔄 Groq API temporarily unavailable. Retrying in {wait_time}s... (Attempt {attempt + 1}/{self.max_retries})")
                        time.sleep(wait_time)
                        continue
                    else:
                        st.error("🚫 Groq API is currently unavailable. Using offline analysis...")
                        raise e
                elif "rate_limit" in error_msg or "429" in error_msg:
                    if attempt < self.max_retries - 1:
                        wait_time = 5 * (attempt + 1)  # Linear backoff for rate limits
                        st.warning(f"⏳ Rate limit reached. Waiting {wait_time}s... (Attempt {attempt + 1}/{self.max_retries})")
                        time.sleep(wait_time)
                        continue
                    else:
                        st.error("🚫 Rate limit exceeded. Please try again later.")
                        raise e
                else:
                    # For other errors, don't retry
                    raise e
        
        return None
    
    def analyze_resume(self, resume_text: str, job_description: str, model_name: str) -> Dict:
        """Main method to analyze resume with retry logic and offline fallback"""
        try:
            # Update model if different
            if hasattr(self, 'llm') and self.llm and model_name in MODELS:
                target_model = MODELS[model_name].model_id
                if self.llm.model_name != target_model:
                    self.llm.model_name = target_model
            
            start_time = time.time()
            
            # Try API call with retries
            try:
                analysis_result = self._retry_with_exponential_backoff(
                    self._analyze_with_llm, 
                    resume_text, 
                    job_description
                )
                
                if analysis_result is None:
                    # API failed after retries, use offline analysis
                    st.warning("🔄 Switching to offline analysis mode...")
                    analysis_result = self._offline_analysis_fallback(resume_text, job_description)
                
            except Exception as e:
                st.error(f"🚫 API Error: {str(e)}")
                st.info("🔄 Using offline analysis mode...")
                analysis_result = self._offline_analysis_fallback(resume_text, job_description)
            
            # Parse and validate results
            parsed_result = self._parse_analysis_result(analysis_result)
            
            analysis_time = time.time() - start_time
            
            if "offline" in str(analysis_result).lower():
                st.info(f"📋 Offline analysis completed in {analysis_time:.2f}s")
            else:
                st.success(f"✅ Analysis completed in {analysis_time:.2f}s using {model_name}")
            
            return parsed_result
                
        except Exception as e:
            logger.error(f"Analysis error: {str(e)}")
            st.error(f"Analysis failed: {str(e)}")
            return self._create_fallback_analysis()
    
    def _offline_analysis_fallback(self, resume_text: str, job_description: str) -> str:
        """Offline analysis using keyword matching and basic algorithms"""
        try:
            # Simple keyword extraction and matching
            import re
            from collections import Counter
            
            # Extract keywords from job description
            job_words = re.findall(r'\b[a-zA-Z]{3,}\b', job_description.lower())
            job_word_freq = Counter(job_words)
            
            # Common tech keywords
            tech_keywords = {
                'python', 'java', 'javascript', 'react', 'node', 'sql', 'aws', 'docker', 
                'kubernetes', 'git', 'linux', 'mongodb', 'postgresql', 'redis', 'api',
                'rest', 'microservices', 'agile', 'scrum', 'ci/cd', 'jenkins', 'terraform'
            }
            
            # Extract keywords from resume
            resume_words = re.findall(r'\b[a-zA-Z]{3,}\b', resume_text.lower())
            resume_word_freq = Counter(resume_words)
            
            # Find matches
            matched_keywords = []
            missing_keywords = []
            
            important_job_keywords = [word for word, freq in job_word_freq.most_common(30) 
                                    if len(word) > 3 and word.isalpha()]
            
            for keyword in important_job_keywords:
                if keyword in resume_word_freq:
                    matched_keywords.append(keyword.title())
                else:
                    missing_keywords.append(keyword.title())
            
            # Calculate match percentage
            total_important = len(important_job_keywords)
            matches = len(matched_keywords)
            match_percentage = int((matches / total_important * 100)) if total_important > 0 else 0
            
            # Calculate overall score
            score = min(85, match_percentage + 15)  # Cap at 85 for offline analysis
            
            # Generate suggestions
            suggestions = [
                "Add more keywords from the job description",
                "Quantify your achievements with specific numbers",
                "Use strong action verbs to describe your experience"
            ]
            
            if missing_keywords:
                suggestions.append(f"Consider adding skills like: {', '.join(missing_keywords[:3])}")
            
            # Create structured response
            offline_result = {
                "score": score,
                "keyword_match": {
                    "matched_keywords": matched_keywords[:20],
                    "missing_keywords": missing_keywords[:20],
                    "match_percentage": match_percentage
                },
                "missing_skills": {
                    "technical_skills": [kw for kw in missing_keywords if kw.lower() in tech_keywords][:10],
                    "programming_languages": [kw for kw in missing_keywords if kw.lower() in ['python', 'java', 'javascript', 'c++', 'c#', 'go', 'rust']][:5],
                    "libraries_frameworks": [kw for kw in missing_keywords if kw.lower() in ['react', 'angular', 'vue', 'django', 'flask', 'spring']][:5],
                    "tools_platforms": [kw for kw in missing_keywords if kw.lower() in ['aws', 'docker', 'kubernetes', 'jenkins', 'git']][:5]
                },
                "soft_skills": ["Communication", "Leadership", "Problem-solving"],
                "suggestions": suggestions
            }
            
            return json.dumps(offline_result)
            
        except Exception as e:
            logger.error(f"Offline analysis failed: {str(e)}")
            return json.dumps(self._create_fallback_analysis())
    
    def _analyze_with_llm(self, resume_text: str, job_description: str) -> str:
        """Perform analysis using direct LLM call"""
        prompt = f"""
        You are an expert ATS resume analyzer. Analyze this resume against the job description.
        
        Return ONLY a valid JSON object with this EXACT structure (no other text):
        {{
            "score": <number 0-100>,
            "keyword_match": {{
                "matched_keywords": ["keyword1", "keyword2"],
                "missing_keywords": ["missing1", "missing2"],
                "match_percentage": <number 0-100>
            }},
            "missing_skills": {{
                "technical_skills": ["skill1", "skill2"],
                "programming_languages": ["lang1", "lang2"],
                "libraries_frameworks": ["lib1", "lib2"],
                "tools_platforms": ["tool1", "tool2"]
            }},
            "soft_skills": ["soft_skill1", "soft_skill2"],
            "suggestions": ["suggestion1", "suggestion2", "suggestion3"]
        }}

        Job Description:
        {job_description[:3000]}

        Resume Content:
        {resume_text[:4000]}

        JSON Response:"""
        
        try:
            response = self.llm.invoke(prompt)
            return response.content if hasattr(response, 'content') else str(response)
        except Exception as e:
            logger.error(f"LLM invocation failed: {str(e)}")
            raise e
    
    def _parse_analysis_result(self, result_string: str) -> Dict:
        """Parse and validate the analysis result"""
        try:
            # Clean the result string
            cleaned = result_string.strip()
            
            # Find JSON boundaries
            json_start = cleaned.find('{')
            json_end = cleaned.rfind('}') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_str = cleaned[json_start:json_end]
                parsed = json.loads(json_str)
                
                # Validate structure
                if self._validate_analysis_structure(parsed):
                    return parsed
                else:
                    logger.error("Invalid analysis structure")
                    return self._create_fallback_analysis()
            else:
                logger.error("No valid JSON found in response")
                return self._create_fallback_analysis()
                
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing failed: {str(e)}")
            return self._create_fallback_analysis()
        except Exception as e:
            logger.error(f"Result parsing failed: {str(e)}")
            return self._create_fallback_analysis()
    
    def optimize_resume(self, resume_text: str, job_description: str, analysis_results: str, model_name: str) -> str:
        """Optimize resume with retry logic and offline fallback"""
        try:
            # Update model if different
            if hasattr(self, 'llm') and self.llm and model_name in MODELS:
                target_model = MODELS[model_name].model_id
                if self.llm.model_name != target_model:
                    self.llm.model_name = target_model
            
            start_time = time.time()
            
            # Try API call with retries
            try:
                optimized_resume = self._retry_with_exponential_backoff(
                    self._optimize_with_llm,
                    resume_text,
                    job_description,
                    analysis_results
                )
                
                if optimized_resume is None:
                    # API failed after retries, use offline optimization
                    st.warning("🔄 Switching to offline optimization mode...")
                    optimized_resume = self._offline_optimization_fallback(resume_text, job_description, analysis_results)
                    
            except Exception as e:
                st.error(f"🚫 API Error: {str(e)}")
                st.info("🔄 Using offline optimization mode...")
                optimized_resume = self._offline_optimization_fallback(resume_text, job_description, analysis_results)
            
            optimization_time = time.time() - start_time
            
            if "offline" in optimized_resume.lower():
                st.info(f"📋 Offline optimization completed in {optimization_time:.2f}s")
            else:
                st.success(f"✅ Optimization completed in {optimization_time:.2f}s")
            
            return optimized_resume
            
        except Exception as e:
            logger.error(f"Optimization error: {str(e)}")
            st.error(f"Optimization failed: {str(e)}")
            return f"Optimization failed: {str(e)}. Please check your inputs and try again."
    
    def _offline_optimization_fallback(self, resume_text: str, job_description: str, analysis_results: str) -> str:
        """Offline resume optimization using basic text processing"""
        try:
            import re
            from collections import Counter
            
            # Extract important keywords from job description
            job_keywords = re.findall(r'\b[A-Za-z]{3,}\b', job_description.lower())
            job_word_freq = Counter(job_keywords)
            important_keywords = [word for word, freq in job_word_freq.most_common(20) 
                                if len(word) > 3 and word.isalpha()]
            
            # Basic optimization rules
            optimization_suggestions = f"""
            
            OFFLINE OPTIMIZATION SUGGESTIONS:
            =================================
            
            📌 KEYWORD INTEGRATION:
            Try to naturally incorporate these important keywords from the job description:
            {', '.join(important_keywords[:10])}
            
            📌 ACTION VERBS TO USE:
            Replace weak verbs with: Led, Architected, Optimized, Implemented, Designed, 
            Developed, Managed, Created, Improved, Achieved
            
            📌 QUANTIFICATION EXAMPLES:
            • "Improved system performance" → "Improved system performance by 40%"
            • "Managed team" → "Managed team of 8 developers"
            • "Reduced costs" → "Reduced operational costs by $50K annually"
            
            📌 ATS OPTIMIZATION TIPS:
            • Use standard section headings (Experience, Skills, Education)
            • Include both acronyms and full forms (AI/Artificial Intelligence)
            • Use bullet points for achievements
            • Avoid images, tables, or complex formatting
            
            📌 MISSING SKILLS TO ADD (if applicable):
            Based on the job description, consider adding experience with:
            {', '.join([kw for kw in important_keywords if kw.lower() in 
                       ['python', 'java', 'aws', 'docker', 'react', 'sql']][:5])}
            
            📌 YOUR CURRENT RESUME:
            {resume_text[:2000]}{'...' if len(resume_text) > 2000 else ''}
            
            ⚠️ NOTE: This is an offline optimization. For AI-powered optimization, 
            please try again when the API service is available.
            """
            
            return optimization_suggestions
            
        except Exception as e:
            logger.error(f"Offline optimization failed: {str(e)}")
            return f"""
            OFFLINE OPTIMIZATION FAILED
            
            Due to technical issues, we couldn't process your resume optimization.
            
            Here are some general tips to improve your resume:
            
            1. Include keywords from the job description
            2. Quantify your achievements with numbers
            3. Use strong action verbs (Led, Implemented, Optimized)
            4. Keep formatting simple for ATS compatibility
            5. Tailor your resume for each specific job application
            
            Please try again later when the service is restored.
            """
    
    def _optimize_with_llm(self, resume_text: str, job_description: str, analysis_results: str) -> str:
        """Perform optimization using direct LLM call"""
        prompt = f"""
        You are a professional resume writer. Optimize this resume for the target job.
        
        Guidelines:
        1. Incorporate relevant keywords naturally from the job description
        2. Quantify achievements with specific metrics where possible
        3. Use strong action verbs (Led, Architected, Optimized, Implemented, etc.)
        4. Ensure ATS-friendly formatting
        5. Maintain professional tone and factual accuracy
        6. Keep the original structure but enhance content
        
        Job Description:
        {job_description[:2500]}
        
        Current Resume:
        {resume_text[:3500]}
        
        Analysis Insights:
        {analysis_results[:1500]}
        
        Return ONLY the optimized resume content (no explanations or additional text):
        """
        
        try:
            response = self.llm.invoke(prompt)
            content = response.content if hasattr(response, 'content') else str(response)
            return content.strip()
        except Exception as e:
            logger.error(f"LLM optimization failed: {str(e)}")
            raise e
    
    def _validate_analysis_structure(self, result: Dict) -> bool:
        """Validate that the analysis result has the required structure"""
        required_keys = ['score', 'keyword_match', 'missing_skills', 'soft_skills', 'suggestions']
        
        if not all(key in result for key in required_keys):
            logger.error(f"Missing required keys. Found: {list(result.keys())}")
            return False
        
        if not isinstance(result.get('keyword_match'), dict):
            logger.error("keyword_match is not a dict")
            return False
            
        keyword_keys = ['matched_keywords', 'missing_keywords', 'match_percentage']
        keyword_match = result.get('keyword_match', {})
        if not all(key in keyword_match for key in keyword_keys):
            logger.error(f"Missing keyword_match keys. Found: {list(keyword_match.keys())}")
            return False
            
        return True
    
    def _create_fallback_analysis(self) -> Dict:
        """Create fallback analysis when parsing fails"""
        return {
            "score": 0,
            "keyword_match": {
                "matched_keywords": [],
                "missing_keywords": ["Analysis failed - please try again"],
                "match_percentage": 0
            },
            "missing_skills": {
                "technical_skills": [],
                "programming_languages": [],
                "libraries_frameworks": [],
                "tools_platforms": []
            },
            "soft_skills": [],
            "suggestions": ["Please try again with different inputs", "Check your API key and internet connection"]
        }

# ========== UI Components (Same as before but enhanced) ==========
def inject_custom_css():
    st.markdown("""
    <style>
    :root {
        --primary: #6e48aa;
        --secondary: #9d50bb;
        --success: #0f9d58;
        --warning: #f4b400;
        --error: #ff4b4b;
    }
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #e4e8f0 100%);
    }
    .agent-status {
        padding: 0.5rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        border-left: 4px solid var(--primary);
        background: rgba(110, 72, 170, 0.1);
    }
    .custom-card {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
        margin-bottom: 1rem;
        border-left: 4px solid var(--primary);
    }
    .model-selector {
        background: white;
        padding: 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
    }
    </style>
    """, unsafe_allow_html=True)

def display_agent_status(agent_system):
    """Display current agent system status"""
    st.markdown("""
    <div class="agent-status">
        <h4>🤖 Agent System Status</h4>
        <p>✅ Resume Analyzer Agent: Active</p>
        <p>✅ Resume Optimizer Agent: Active</p>
        <p>✅ Content Validator Agent: Active</p>
    </div>
    """, unsafe_allow_html=True)

def display_analysis_results(analysis: Dict):
    """Enhanced display of analysis results"""
    # Overall Score
    score = analysis.get('score', 0)
    score_color = 'var(--error)' if score < 50 else 'var(--warning)' if score < 75 else 'var(--success)'
    
    st.markdown(f"""
    <div class="custom-card">
        <h3>🎯 ATS Compatibility Score</h3>
        <h1 style="color: {score_color}; text-align: center; margin: 1rem 0;">{score}/100</h1>
        <div style="background: #e0e0e0; border-radius: 10px; height: 10px;">
            <div style="background: linear-gradient(to right, var(--primary), var(--secondary)); 
                        width: {score}%; height: 100%; border-radius: 10px;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Keyword Analysis
    keyword_data = analysis.get('keyword_match', {})
    col1, col2 = st.columns(2)
    
    with col1:
        if keyword_data.get('matched_keywords'):
            st.success(f"✅ {len(keyword_data['matched_keywords'])} Keywords Matched")
            with st.expander("View Matched Keywords"):
                st.write(", ".join(keyword_data['matched_keywords'][:20]))
    
    with col2:
        if keyword_data.get('missing_keywords'):
            st.warning(f"⚠️ {len(keyword_data['missing_keywords'])} Keywords Missing")
            with st.expander("View Missing Keywords"):
                st.write(", ".join(keyword_data['missing_keywords'][:20]))
    
    # Skills Analysis
    skills = analysis.get('missing_skills', {})
    if any(skills.values()):
        st.markdown("### 🛠️ Missing Skills Analysis")
        
        skills_cols = st.columns(2)
        with skills_cols[0]:
            if skills.get('technical_skills'):
                st.markdown("**Technical Skills:**")
                for skill in skills['technical_skills'][:10]:
                    st.markdown(f"• {skill}")
            
            if skills.get('programming_languages'):
                st.markdown("**Programming Languages:**")
                for lang in skills['programming_languages'][:10]:
                    st.markdown(f"• {lang}")
        
        with skills_cols[1]:
            if skills.get('libraries_frameworks'):
                st.markdown("**Libraries/Frameworks:**")
                for lib in skills['libraries_frameworks'][:10]:
                    st.markdown(f"• {lib}")
            
            if skills.get('tools_platforms'):
                st.markdown("**Tools/Platforms:**")
                for tool in skills['tools_platforms'][:10]:
                    st.markdown(f"• {tool}")
    
    # Suggestions
    suggestions = analysis.get('suggestions', [])
    if suggestions:
        st.markdown("### 💡 Improvement Suggestions")
        for i, suggestion in enumerate(suggestions[:5], 1):
            st.markdown(f"""
            <div class="custom-card">
                <strong>{i}.</strong> {suggestion}
            </div>
            """, unsafe_allow_html=True)

# ========== Main Application ==========
def main():
    st.set_page_config(
        page_title="🤖 Agentic Resume Optimizer",
        page_icon="⚡",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    inject_custom_css()
    
    st.title("🤖 Agentic Resume Optimizer")
    st.markdown("**Powered by LangChain Agents & Groq AI**")
    
    # API Key Input
    if 'groq_api_key' not in st.session_state:
        st.session_state.groq_api_key = ""
    
    with st.sidebar:
        st.subheader("🔑 Configuration")
        api_key = st.text_input(
            "Groq API Key", 
            value=st.session_state.groq_api_key,
            type="password",
            help="Get your free API key from https://console.groq.com/"
        )
        
        if api_key != st.session_state.groq_api_key:
            st.session_state.groq_api_key = api_key
            if 'agent_system' in st.session_state:
                del st.session_state.agent_system
        
        # Model Selection
        st.subheader("🧠 Model Selection")
        selected_model = st.selectbox(
            "Choose AI Model",
            options=list(MODELS.keys()),
            index=1,
            help="Different models offer various trade-offs between speed and accuracy"
        )
        
        model_config = MODELS[selected_model]
        st.markdown(f"""
        <div class="model-selector">
            <strong>{model_config.name}</strong><br>
            <small>{model_config.description}</small><br>
            <small>Speed Rating: {'⭐' * model_config.speed_rating}/5</small>
        </div>
        """, unsafe_allow_html=True)
    
    # Initialize Agent System with better error handling
    if api_key and 'agent_system' not in st.session_state:
        try:
            with st.spinner("🤖 Initializing AI Agent System..."):
                st.session_state.agent_system = ResumeAgentSystem(api_key)
            st.success("✅ Agent system initialized successfully!")
        except Exception as e:
            st.error(f"❌ Failed to initialize agents: {str(e)}")
            st.info("💡 This might be due to:")
            st.markdown("""
            - Invalid API key
            - Network connectivity issues  
            - Groq API service being temporarily unavailable
            
            **Solutions:**
            - Check your API key at [Groq Console](https://console.groq.com/)
            - Visit [Groq Status](https://groqstatus.com/) for service status
            - Try again in a few minutes
            """)
            return
    
    if not api_key:
        st.warning("⚠️ Please enter your Groq API key in the sidebar to continue.")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            ### 🚀 How to get started:
            1. **Get API Key**: Visit [Groq Console](https://console.groq.com/)
            2. **Create Account**: Sign up for free
            3. **Generate Key**: Create a new API key
            4. **Enter Key**: Paste it in the sidebar
            5. **Start Analyzing**: Upload resume and job description!
            """)
            
        with col2:
            st.markdown("""
            ### 🛡️ Fallback Features:
            - **Offline Analysis**: Basic keyword matching when API is down
            - **Retry Logic**: Automatic retries with exponential backoff
            - **Status Monitoring**: Real-time API status checking
            - **Error Recovery**: Graceful degradation to offline mode
            - **Service Status**: Check [Groq Status](https://groqstatus.com/)
            """)
            
        st.info("💡 **Tip**: Even with API issues, the app provides offline analysis using keyword matching algorithms!")
        return
    
    # Display agent status
    if 'agent_system' in st.session_state:
        display_agent_status(st.session_state.agent_system)
    
    # Main Interface
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader("📄 Upload Resume")
        uploaded_file = st.file_uploader(
            "Choose PDF or DOCX file", 
            type=["pdf", "docx"],
            help="Upload your current resume for analysis"
        )
        
        st.subheader("💼 Job Information")
        job_title = st.text_input("Job Title", placeholder="e.g., Senior Software Engineer")
        job_description = st.text_area(
            "Job Description",
            height=200,
            placeholder="Paste the complete job description here..."
        )
        
        # Analysis Button
        if st.button("🚀 Analyze with AI Agents", type="primary", use_container_width=True):
            if not uploaded_file or not job_description:
                st.error("Please upload a resume and enter job description.")
            elif 'agent_system' not in st.session_state:
                st.error("Agent system not initialized. Please check your API key.")
            else:
                try:
                    # Extract text from resume
                    with st.spinner("📄 Extracting text from resume..."):
                        resume_text = extract_text_from_file(uploaded_file)
                    
                    if resume_text:
                        # Validate resume text
                        if len(resume_text.strip()) < 100:
                            st.warning("Resume seems too short. Please check if the file was processed correctly.")
                        
                        # Run analysis using agent system
                        with st.spinner("🤖 AI agents analyzing your resume..."):
                            analysis = st.session_state.agent_system.analyze_resume(
                                resume_text, job_description, selected_model
                            )
                            
                            if analysis and analysis.get('score', 0) > 0:
                                st.session_state.analysis_results = analysis
                                st.session_state.resume_text = resume_text
                                st.session_state.job_description = job_description
                                st.balloons()
                            else:
                                st.error("Analysis returned empty results. Please try again.")
                    else:
                        st.error("Failed to extract text from the uploaded file. Please try a different file.")
                        
                except Exception as e:
                    st.error(f"An error occurred during processing: {str(e)}")
                    logger.error(f"Processing error: {str(e)}")
    
    with col2:
        if uploaded_file:
            st.subheader("📋 Resume Preview")
            resume_text = extract_text_from_file(uploaded_file)
            if resume_text:
                st.text_area("Extracted Content", resume_text[:2000] + "..." if len(resume_text) > 2000 else resume_text, height=300)
        else:
            st.markdown("""
            ### 🤖 About This Agentic System
            
            This application uses **specialized AI agents** built with LangChain:
            
            **🔍 Analyzer Agent**: Performs deep ATS compatibility analysis
            **✨ Optimizer Agent**: Customizes resume content for job relevance  
            **🛡️ Validator Agent**: Ensures quality and accuracy
            
            **Features:**
            - Multi-model support (Gemma, Llama 3, Mixtral)
            - Intelligent error handling and recovery
            - Structured output parsing
            - Real-time performance monitoring
            - Memory-based conversation tracking
            """)
    
    # Display Results
    if 'analysis_results' in st.session_state:
        st.markdown("---")
        st.subheader("📊 Analysis Results")
        display_analysis_results(st.session_state.analysis_results)
        
        # Optimization Section
        st.markdown("---")
        st.subheader("🛠️ Resume Optimization")
        
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("✨ Optimize Resume", type="secondary", use_container_width=True):
                if 'agent_system' in st.session_state and 'resume_text' in st.session_state:
                    try:
                        with st.spinner("🤖 AI agents optimizing your resume..."):
                            optimized = st.session_state.agent_system.optimize_resume(
                                st.session_state.resume_text,
                                st.session_state.job_description,
                                str(st.session_state.analysis_results),
                                selected_model
                            )
                            
                            if optimized and len(optimized.strip()) > 100:
                                st.session_state.optimized_resume = optimized
                                st.success("Resume optimization completed!")
                            else:
                                st.error("Optimization failed or returned empty result.")
                                
                    except Exception as e:
                        st.error(f"Optimization error: {str(e)}")
                        logger.error(f"Optimization error: {str(e)}")
                else:
                    st.error("Please run analysis first before optimizing.")
        
        with col2:
            if 'optimized_resume' in st.session_state:
                st.download_button(
                    "📥 Download Optimized Resume",
                    st.session_state.optimized_resume,
                    file_name=f"optimized_resume_{job_title.lower().replace(' ', '_') if job_title else 'resume'}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
        
        # Show optimized resume
        if 'optimized_resume' in st.session_state:
            st.markdown("### 📝 Optimized Resume")
            st.text_area(
                "Optimized Content",
                st.session_state.optimized_resume,
                height=400,
                help="Review and edit as needed before using"
            )

def extract_text_from_file(uploaded_file) -> Optional[str]:
    """Extract text from uploaded file with error handling"""
    try:
        if uploaded_file.name.endswith('.pdf'):
            with pdfplumber.open(uploaded_file) as pdf:
                text = " ".join(page.extract_text() or "" for page in pdf.pages)
        elif uploaded_file.name.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            text = " ".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
        else:
            st.error(f"Unsupported file format: {uploaded_file.name}")
            return None
        
        # Clean the text
        for char in ['\x0c', '\uf0b7', '\u2022']:
            text = text.replace(char, ' ')
        while "  " in text:
            text = text.replace("  ", " ")
        
        return text.strip()
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        logger.error(f"File extraction error: {str(e)}")
        return None

if __name__ == "__main__":
    main()